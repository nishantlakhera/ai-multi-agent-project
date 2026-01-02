from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
import re
from typing import Dict, List, Optional

from docx import Document


@dataclass
class ConvertedCase:
    scenario: str
    steps: List[str]
    test_data: Dict[str, str]
    expected: str
    tags: List[str]


def _sanitize_scenario(text: str) -> str:
    cleaned = re.sub(r"[_\-]+", " ", text).strip()
    return re.sub(r"\s+", " ", cleaned)


def _derive_scenario(prefix: Optional[str], filename: str, index: int) -> str:
    base = Path(filename).stem or f"recording-{index}"
    name = _sanitize_scenario(base)
    if prefix:
        return f"{prefix} - {name}"
    return name


def _selector_from_locator(locator: str) -> Optional[str]:
    if not locator:
        return None
    return locator


def _locator_from_line(line: str) -> Optional[str]:
    patterns = [
        (r"getByRole\(\s*['\"](\w+)['\"]\s*,\s*\{[^}]*name:\s*['\"]([^'\"]+)['\"][^}]*\}\s*\)", "role"),
        (r"get_by_role\(\s*['\"](\w+)['\"]\s*,\s*name=['\"]([^'\"]+)['\"]\s*\)", "role"),
        (r"getByLabel\(\s*['\"]([^'\"]+)['\"]\s*\)", "label"),
        (r"get_by_label\(\s*['\"]([^'\"]+)['\"]\s*\)", "label"),
        (r"getByPlaceholder\(\s*['\"]([^'\"]+)['\"]\s*\)", "placeholder"),
        (r"get_by_placeholder\(\s*['\"]([^'\"]+)['\"]\s*\)", "placeholder"),
        (r"getByText\(\s*['\"]([^'\"]+)['\"]\s*\)", "text"),
        (r"get_by_text\(\s*['\"]([^'\"]+)['\"]\s*\)", "text"),
        (r"locator\(\s*['\"]([^'\"]+)['\"]\s*\)", "css"),
    ]

    for pattern, kind in patterns:
        match = re.search(pattern, line)
        if not match:
            continue
        if kind == "role":
            role, name = match.group(1), match.group(2)
            return f'role={role} name="{name}"'
        value = match.group(1)
        return f'{kind}="{value}"'
    return None


def _value_from_fill(line: str, constants: Dict[str, str]) -> Optional[str]:
    match = re.search(r"\.fill\(\s*['\"]([^'\"]*)['\"]\s*\)", line)
    if match:
        return match.group(1)
    match = re.search(r"\.fill\(\s*([A-Z_][A-Z0-9_]*)\s*\)", line)
    if match:
        key = match.group(1)
        return constants.get(key)
    return None


def _value_from_select(line: str) -> Optional[str]:
    match = re.search(r"\.select(?:_option|Option)\(\s*['\"]([^'\"]+)['\"]\s*\)", line)
    if match:
        return match.group(1)
    match = re.search(r"\.select(?:_option|Option)\(\s*\{[^}]*label:\s*['\"]([^'\"]+)['\"]", line)
    if match:
        return match.group(1)
    return None


def _normalize_key(raw: str, index: int) -> str:
    key = raw.strip().lower().rstrip(":")
    key = re.sub(r"[^a-z0-9]+", "_", key).strip("_")
    return key or f"field_{index}"


def _key_from_selector(selector: str, index: int) -> str:
    name_match = re.search(r'name=\"([^\"]+)\"', selector)
    if name_match:
        return _normalize_key(name_match.group(1), index)
    label_match = re.search(r'label=\"([^\"]+)\"', selector)
    if label_match:
        return _normalize_key(label_match.group(1), index)
    placeholder_match = re.search(r'placeholder=\"([^\"]+)\"', selector)
    if placeholder_match:
        return _normalize_key(placeholder_match.group(1), index)
    if selector.startswith("css="):
        id_match = re.search(r"#([A-Za-z0-9_-]+)", selector)
        if id_match:
            return _normalize_key(id_match.group(1), index)
    return _normalize_key(selector, index)


def _extract_constants(script: str) -> Dict[str, str]:
    constants: Dict[str, str] = {}
    for line in script.splitlines():
        match = re.match(r'^\s*([A-Z_][A-Z0-9_]*)\s*=\s*[\'"]([^\'"]+)[\'"]\s*$', line)
        if match:
            constants[match.group(1)] = match.group(2)
    return constants


def _parse_recording(script: str) -> tuple[List[str], Dict[str, str]]:
    steps: List[str] = []
    test_data: Dict[str, str] = {}
    lines = script.splitlines()
    constants = _extract_constants(script)

    for line in lines:
        if ".goto(" in line:
            match = re.search(r"\.goto\(\s*['\"]([^'\"]+)['\"]\s*\)", line)
            if match:
                steps.append(f"Goto {match.group(1)}")
            continue

        if ".click(" in line:
            locator = _locator_from_line(line)
            if locator:
                steps.append(f"Click on element (selector: {_selector_from_locator(locator)})")
            continue

        if ".fill(" in line:
            locator = _locator_from_line(line)
            value = _value_from_fill(line, constants)
            if locator:
                steps.append(f"Enter value (selector: {_selector_from_locator(locator)})")
            if locator and value is not None:
                key = _key_from_selector(locator, len(test_data) + 1)
                test_data.setdefault(key, value)
            continue

        if ".select" in line:
            locator = _locator_from_line(line)
            value = _value_from_select(line)
            if locator and value:
                steps.append(f"Select {value} from dropdown (selector: {_selector_from_locator(locator)})")
            continue

    return steps, test_data


def _ensure_doc(doc_path: Path) -> Document:
    if doc_path.exists():
        return Document(doc_path)

    doc = Document()
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["Scenario", "Test Steps", "Test Data", "Expected", "Tags"]
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    return doc


def append_cases_to_docx(doc_path: str, cases: List[ConvertedCase]) -> int:
    path = Path(doc_path)
    doc = _ensure_doc(path)
    table = doc.tables[0]

    for case in cases:
        row = table.add_row()
        row.cells[0].text = case.scenario
        row.cells[1].text = "\n".join(case.steps)
        row.cells[2].text = "\n".join([f"{k}: {v}" for k, v in case.test_data.items()])
        row.cells[3].text = case.expected
        row.cells[4].text = ", ".join(case.tags)

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    return len(cases)


def convert_recordings(
    recordings: List[tuple[str, str]],
    doc_path: str,
    scenario_prefix: Optional[str] = None,
    expected: Optional[str] = None,
    tags: Optional[List[str]] = None,
) -> int:
    cases: List[ConvertedCase] = []
    tags = tags or []
    expected = expected or "Expected result should be displayed."

    for idx, (filename, content) in enumerate(recordings, 1):
        steps, test_data = _parse_recording(content)
        scenario = _derive_scenario(scenario_prefix, filename, idx)
        cases.append(
            ConvertedCase(
                scenario=scenario,
                steps=steps or ["<no steps parsed>"],
                test_data=test_data,
                expected=expected,
                tags=tags,
            )
        )

    return append_cases_to_docx(doc_path, cases)
