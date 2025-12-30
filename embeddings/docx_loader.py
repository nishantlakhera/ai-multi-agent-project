from pathlib import Path
from typing import List, Dict
import uuid
from docx import Document
from embeddings.document_loader import chunk_text

def _extract_docx_text(path: Path) -> str:
    doc = Document(path)
    parts: List[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)

def load_docx_files(folder: str, chunk_documents: bool = True) -> List[Dict]:
    """
    Load .docx files from a folder, extract text, and optionally chunk them.
    """
    docs: List[Dict] = []
    folder_path = Path(folder)
    if not folder_path.exists():
        print(f"Warning: Folder {folder} does not exist")
        return []

    docx_files = list(folder_path.rglob("*.docx"))
    if not docx_files:
        print(f"Warning: No .docx files found in {folder}")
        return []

    for path in docx_files:
        try:
            text = _extract_docx_text(path)
            if not text.strip():
                continue

            if chunk_documents and len(text) > 1000:
                chunks = chunk_text(text, chunk_size=1000, overlap=200)
                for i, chunk in enumerate(chunks):
                    if chunk.strip():
                        docs.append({
                            "id": str(uuid.uuid4()),
                            "text": chunk,
                            "meta": {
                                "path": str(path),
                                "source": "docx_file",
                                "filename": path.name,
                                "doc_type": "testcase",
                                "chunk_index": i,
                                "total_chunks": len(chunks),
                            },
                        })
            else:
                docs.append({
                    "id": str(uuid.uuid4()),
                    "text": text,
                    "meta": {
                        "path": str(path),
                        "source": "docx_file",
                        "filename": path.name,
                        "doc_type": "testcase",
                    },
                })
        except Exception as e:
            print(f"Error loading {path}: {e}")
            continue

    print(f"Loaded {len(docs)} document chunks from {len(docx_files)} files")
    return docs
